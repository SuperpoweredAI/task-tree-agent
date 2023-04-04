from docx import Document as DocxDocument # pip install python-docx

def format_txt(document):
    formatted = f"Title: {document.title}\n\nTable of Contents:\n"
    for index, title in enumerate(document.table_of_contents):
        formatted += f"{index + 1}. {title}\n"
    formatted += "\n"

    for section in document.sections:
        formatted += format_section_txt(section)

    return formatted

def format_section_txt(section, level=1):
    formatted = f"{'#' * level} {section.title}\n\n"
    for element in section.elements:
        formatted += f"{element.content}\n\n"

    return formatted

def format_markdown(document):
    formatted = f"# {document.title}\n\n## Table of Contents\n"
    for index, title in enumerate(document.table_of_contents):
        formatted += f"{index + 1}. [{title}](#{title.lower().replace(' ', '-')})\n"
    formatted += "\n"

    for section in document.sections:
        formatted += format_section_markdown(section)

    return formatted

def format_section_markdown(section, level=1):
    formatted = f"{'#' * (level + 1)} {section.title}\n\n"
    for element in section.elements:
        formatted += f"{element.content}\n\n"

    return formatted

def format_docx(document, output_file):
    docx = DocxDocument()

    # Add title and author
    docx.add_heading(document.title, level=1)
    docx.add_paragraph(f"Author: {document.author}")

    # Add table of contents
    docx.add_heading("Table of Contents", level=2)
    for index, title in enumerate(document.table_of_contents):
        docx.add_paragraph(f"{index + 1}. {title}")

    # Add sections
    for section in document.sections:
        format_section_docx(docx, section)

    # Save the docx file
    docx.save(output_file)

def format_section_docx(docx, section, level=2):
    docx.add_heading(section.title, level=level)
    for element in section.elements:
        if element.element_type == 'text':
            docx.add_paragraph(element.content)
        elif element.element_type == 'image':
            docx.add_paragraph(f"[Image: {element.content}]")
        elif element.element_type == 'table':
            docx.add_paragraph(f"[Table: {element.content}]")
        elif element.element_type == 'list':
            docx.add_paragraph(f"[List: {element.content}]")
        elif element.element_type == 'code':
            docx.add_paragraph(f"[Code: {element.content}]")
        else:
            docx.add_paragraph(f"[Unknown element type: {element.element_type}]")

    for subsection in section.subsections:
        format_section_docx(docx, subsection, level=level + 1)


import pickle
import sys
import os

sys.path.append(os.path.join(os.path.dirname(__file__), "../elegant_infinite_agent"))

from elegant_infinite_agent import Agent

with open("technology_and_society.pkl", "rb") as f:
    agent = pickle.load(f)

document = agent.sdf_document
md_document = format_markdown(document)

# save the markdown file
with open("technology_and_society.md", "w") as f:
    f.write(md_document)